"""
Enhanced Resume Parser with Advanced Text Extraction and Entity Recognition
Includes support for PDF, DOC/DOCX, images (OCR), and plain text
"""

import re
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import os
from pathlib import Path

# Document parsing libraries
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
    from docx import Document
except ImportError:
    docx = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# NLP libraries
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except:
    spacy = None
    nlp = None

# Date parsing
from dateutil import parser as date_parser
from dateutil.relativedelta import relativedelta

# Fuzzy matching
try:
    from rapidfuzz import fuzz, process
except ImportError:
    fuzz = None
    process = None

logger = logging.getLogger(__name__)

class EnhancedResumeParser:
    """Advanced resume parser with multiple extraction methods and NLP capabilities"""
    
    def __init__(self):
        self.skills_database = self._load_skills_database()
        self.job_titles_database = self._load_job_titles_database()
        
    def _load_skills_database(self) -> List[str]:
        """Load comprehensive skills database"""
        # Common technical skills
        return [
            # Programming Languages
            "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Ruby", "Go", "Rust", "Swift",
            "Kotlin", "PHP", "R", "MATLAB", "Scala", "Perl", "Objective-C", "Visual Basic", "Dart",
            
            # Web Technologies
            "HTML", "CSS", "React", "Angular", "Vue.js", "Node.js", "Express.js", "Django", "Flask",
            "Spring Boot", "ASP.NET", "Ruby on Rails", "Laravel", "Symfony", "Next.js", "Nuxt.js",
            
            # Databases
            "MySQL", "PostgreSQL", "MongoDB", "Oracle", "SQL Server", "Redis", "Cassandra", "DynamoDB",
            "Elasticsearch", "Neo4j", "CouchDB", "Firebase", "Supabase",
            
            # Cloud & DevOps
            "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes", "Jenkins", "GitLab CI", "GitHub Actions",
            "Terraform", "Ansible", "CloudFormation", "Helm", "ArgoCD", "Prometheus", "Grafana",
            
            # Data Science & ML
            "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "Pandas",
            "NumPy", "Keras", "Computer Vision", "NLP", "OpenCV", "YOLO", "Transformers",
            
            # Mobile Development
            "React Native", "Flutter", "Android", "iOS", "Xamarin", "Ionic", "SwiftUI", "Jetpack Compose",
            
            # Other Technical Skills
            "Git", "REST API", "GraphQL", "Microservices", "Agile", "Scrum", "JIRA", "Linux", "Windows Server",
            "Networking", "Security", "Blockchain", "IoT", "AR/VR", "Game Development", "Unity", "Unreal Engine"
        ]
    
    def _load_job_titles_database(self) -> List[str]:
        """Load common job titles"""
        return [
            "Software Engineer", "Senior Software Engineer", "Lead Software Engineer", "Principal Engineer",
            "Full Stack Developer", "Frontend Developer", "Backend Developer", "DevOps Engineer",
            "Data Scientist", "Data Engineer", "Machine Learning Engineer", "AI Engineer",
            "Cloud Architect", "Solutions Architect", "Technical Lead", "Engineering Manager",
            "Product Manager", "Project Manager", "Scrum Master", "Business Analyst",
            "QA Engineer", "Test Engineer", "Security Engineer", "Database Administrator",
            "Mobile Developer", "iOS Developer", "Android Developer", "Game Developer",
            "UI/UX Designer", "Technical Writer", "Site Reliability Engineer", "Platform Engineer"
        ]
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self._extract_from_image(file_path)
        elif file_ext in ['.txt', '.text']:
            return self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Method 1: PyMuPDF (fastest and most reliable)
        if fitz:
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text()
                doc.close()
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: pdfplumber (better for tables)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 3: OCR fallback
        if Image and pytesseract:
            try:
                # Convert PDF to images and OCR
                if fitz:
                    doc = fitz.open(file_path)
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        image = Image.open(BytesIO(img_data))
                        text += pytesseract.image_to_string(image) + "\n"
                    doc.close()
            except Exception as e:
                logger.warning(f"OCR extraction failed: {e}")
        
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        if not docx:
            raise ImportError("python-docx is required for DOCX extraction")
        
        doc = Document(file_path)
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n".join(text)
    
    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        if not Image or not pytesseract:
            raise ImportError("PIL and pytesseract are required for image extraction")
        
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use utf-8 with error handling
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def parse_resume(self, text: str, filename: str = "") -> Dict[str, Any]:
        """Parse resume text and extract structured information"""
        # Clean text
        text = self._clean_text(text)
        
        # Extract all components
        result = {
            "filename": filename,
            "contact_info": self._extract_contact_info(text),
            "name": self._extract_name(text),
            "skills": self._extract_skills(text),
            "experience": self._extract_experience(text),
            "education": self._extract_education(text),
            "certifications": self._extract_certifications(text),
            "projects": self._extract_projects(text),
            "languages": self._extract_languages(text),
            "summary": self._extract_summary(text),
            "total_experience_years": self._calculate_total_experience(text),
            "suggested_job_title": self._suggest_job_title(text),
            "parsed_at": datetime.now().isoformat()
        }
        
        return result
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s@.\-+()/:,|]', ' ', text)
        # Normalize case for better matching
        return text.strip()
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information"""
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone
        phone_pattern = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
        phones = re.findall(phone_pattern, text)
        if phones:
            # Clean phone number
            phone = re.sub(r'[^\d+]', '', phones[0])
            if len(phone) >= 10:
                contact['phone'] = phones[0]
        
        # LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
        linkedin = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact['linkedin'] = linkedin.group(0)
        
        # GitHub
        github_pattern = r'github\.com/[\w\-]+'
        github = re.search(github_pattern, text, re.IGNORECASE)
        if github:
            contact['github'] = github.group(0)
        
        return contact
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name using NLP and patterns"""
        # Try NLP first
        if nlp:
            doc = nlp(text[:500])  # Process first 500 chars
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text
        
        # Fallback to pattern matching
        # Usually name appears at the beginning
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            # Skip if line contains common resume keywords
            skip_keywords = ['resume', 'cv', 'curriculum', 'contact', 'email', 'phone']
            if any(keyword in line.lower() for keyword in skip_keywords):
                continue
            
            # Check if line looks like a name (2-4 words, title case)
            words = line.strip().split()
            if 2 <= len(words) <= 4:
                if all(word[0].isupper() for word in words if word):
                    return line.strip()
        
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills using database and NLP"""
        skills = set()
        text_lower = text.lower()
        
        # Match against skills database
        for skill in self.skills_database:
            # Case-insensitive word boundary matching
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                skills.add(skill)
        
        # Extract from skills section
        skills_section_pattern = r'(?:skills|technical skills|core competencies)[\s:]*(.+?)(?:\n\n|\n[A-Z]|$)'
        skills_match = re.search(skills_section_pattern, text, re.IGNORECASE | re.DOTALL)
        if skills_match:
            skills_text = skills_match.group(1)
            # Extract comma or pipe separated skills
            potential_skills = re.split(r'[,|•·]', skills_text)
            for skill in potential_skills:
                skill = skill.strip()
                if 2 < len(skill) < 30:  # Reasonable skill length
                    skills.add(skill)
        
        return list(skills)[:20]  # Limit to 20 skills
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience with advanced parsing"""
        experiences = []
        
        # Common section headers
        exp_headers = [
            'experience', 'work experience', 'professional experience', 
            'employment history', 'work history', 'career history'
        ]
        
        # Find experience section
        exp_section = None
        for header in exp_headers:
            pattern = rf'(?i)\b{header}\b[\s:]*(.+?)(?=\b(?:education|skills|projects|certifications|references)\b|$)'
            match = re.search(pattern, text, re.DOTALL)
            if match:
                exp_section = match.group(1)
                break
        
        if not exp_section:
            exp_section = text  # Process entire text if section not found
        
        # Extract individual experiences
        # Pattern for job entries (Company - Title - Dates)
        job_pattern = r'([A-Z][^•\n]{10,100})\s*[\-–|]\s*([^•\n]{5,100})\s*[\-–|]\s*([^•\n]{5,50})'
        
        matches = re.finditer(job_pattern, exp_section)
        
        for match in matches:
            company = match.group(1).strip()
            title = match.group(2).strip()
            date_str = match.group(3).strip()
            
            # Parse dates
            start_date, end_date = self._parse_date_range(date_str)
            
            # Extract responsibilities (bullet points after the job entry)
            resp_start = match.end()
            next_job = re.search(job_pattern, exp_section[resp_start:])
            resp_end = resp_start + next_job.start() if next_job else len(exp_section)
            
            responsibilities_text = exp_section[resp_start:resp_end]
            responsibilities = self._extract_bullet_points(responsibilities_text)
            
            experiences.append({
                "company": company,
                "title": title,
                "start_date": start_date,
                "end_date": end_date,
                "duration": self._calculate_duration(start_date, end_date),
                "responsibilities": responsibilities[:10]  # Limit to 10
            })
        
        # Sort by date (most recent first)
        experiences.sort(key=lambda x: x['start_date'] or '', reverse=True)
        
        return experiences[:2]  # Return two most recent
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        
        # Education section pattern
        edu_pattern = r'(?i)(?:education|academic|qualification)[\s:]*(.+?)(?=\b(?:experience|skills|projects|certifications)\b|$)'
        edu_match = re.search(edu_pattern, text, re.DOTALL)
        
        if edu_match:
            edu_text = edu_match.group(1)
            
            # Common degree patterns
            degree_patterns = [
                r'(Bachelor(?:\'s)?|B\.?S\.?|B\.?A\.?|Master(?:\'s)?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|B\.?Tech|M\.?Tech)[^,\n]*',
                r'(?:in|of)\s+([A-Za-z\s]+?)(?:from|at|\n|,)',
            ]
            
            # Extract degrees
            degrees = re.findall(degree_patterns[0], edu_text, re.IGNORECASE)
            
            for degree in degrees:
                # Try to find associated institution
                inst_pattern = r'{}\s*(?:from|at|,)\s*([A-Za-z\s]+?)(?:\n|,|$)'.format(re.escape(degree))
                inst_match = re.search(inst_pattern, edu_text, re.IGNORECASE)
                
                education.append({
                    "degree": degree.strip(),
                    "institution": inst_match.group(1).strip() if inst_match else "",
                    "field": self._extract_field_of_study(degree)
                })
        
        return education
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract professional certifications"""
        certifications = []
        
        # Certification section
        cert_pattern = r'(?i)(?:certification|certified|certificate)s?[\s:]*(.+?)(?=\b(?:experience|skills|education|projects)\b|$)'
        cert_match = re.search(cert_pattern, text, re.DOTALL)
        
        if cert_match:
            cert_text = cert_match.group(1)
            # Split by common delimiters
            certs = re.split(r'[•\n,;]', cert_text)
            
            for cert in certs:
                cert = cert.strip()
                if 5 < len(cert) < 100:  # Reasonable cert length
                    certifications.append(cert)
        
        # Also look for common certification patterns
        common_certs = [
            r'(AWS\s+(?:Solutions\s+)?(?:Certified\s+)?[\w\s]+)',
            r'((?:Microsoft\s+)?(?:Azure|AZ-\d+)[\w\s]+)',
            r'((?:Google\s+)?(?:Cloud|GCP)[\w\s]+)',
            r'(PMP|CISSP|CCNA|CCNP|ITIL|Scrum\s+Master|Six\s+Sigma)',
        ]
        
        for pattern in common_certs:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)
        
        return list(set(certifications))[:10]  # Remove duplicates, limit to 10
    
    def _extract_projects(self, text: str) -> List[Dict[str, str]]:
        """Extract project information"""
        projects = []
        
        # Project section
        proj_pattern = r'(?i)(?:projects?)[\s:]*(.+?)(?=\b(?:experience|skills|education|certifications)\b|$)'
        proj_match = re.search(proj_pattern, text, re.DOTALL)
        
        if proj_match:
            proj_text = proj_match.group(1)
            
            # Split by project indicators
            project_splits = re.split(r'(?:^|\n)(?=(?:[A-Z][^:]+:|•|\d+\.))', proj_text)
            
            for proj in project_splits:
                if not proj.strip():
                    continue
                
                # Extract project name (usually first line or before colon)
                lines = proj.strip().split('\n')
                if lines:
                    name = lines[0].strip(' •:')
                    description = ' '.join(lines[1:]) if len(lines) > 1 else ""
                    
                    # Extract technologies used
                    tech_pattern = r'(?:technologies?|tech\s*stack|built\s*with|using)[\s:]*([^.\n]+)'
                    tech_match = re.search(tech_pattern, proj, re.IGNORECASE)
                    technologies = tech_match.group(1) if tech_match else ""
                    
                    projects.append({
                        "name": name[:100],  # Limit length
                        "description": description[:300],
                        "technologies": technologies
                    })
        
        return projects[:5]  # Limit to 5 projects
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract spoken languages"""
        languages = []
        
        # Language section
        lang_pattern = r'(?i)(?:languages?)[\s:]*(.+?)(?=\b(?:experience|skills|education|projects)\b|\n\n|$)'
        lang_match = re.search(lang_pattern, text)
        
        if lang_match:
            lang_text = lang_match.group(1)
            # Common language names
            common_languages = [
                'English', 'Spanish', 'French', 'German', 'Chinese', 'Japanese',
                'Korean', 'Arabic', 'Hindi', 'Portuguese', 'Russian', 'Italian'
            ]
            
            for lang in common_languages:
                if lang.lower() in lang_text.lower():
                    languages.append(lang)
        
        return languages
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary/objective"""
        # Summary section patterns
        summary_patterns = [
            r'(?i)(?:summary|objective|profile|about)[\s:]*(.+?)(?=\b(?:experience|skills|education)\b|\n\n)',
            r'^([^•\n]{50,300})(?=\n)'  # First paragraph if no section header
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                summary = match.group(1).strip()
                # Clean up
                summary = re.sub(r'\s+', ' ', summary)
                return summary[:500]  # Limit length
        
        return ""
    
    def _calculate_total_experience(self, text: str) -> float:
        """Calculate total years of experience"""
        experiences = self._extract_experience(text)
        
        if not experiences:
            # Try to extract from summary/objective
            years_pattern = r'(\d+)\+?\s*years?\s*(?:of\s*)?experience'
            match = re.search(years_pattern, text, re.IGNORECASE)
            if match:
                return float(match.group(1))
            return 0.0
        
        # Calculate from work experiences
        total_months = 0
        for exp in experiences:
            if exp['duration']:
                # Parse duration string (e.g., "2 years 3 months")
                years_match = re.search(r'(\d+)\s*year', exp['duration'])
                months_match = re.search(r'(\d+)\s*month', exp['duration'])
                
                years = int(years_match.group(1)) if years_match else 0
                months = int(months_match.group(1)) if months_match else 0
                
                total_months += (years * 12) + months
        
        return round(total_months / 12, 1)
    
    def _suggest_job_title(self, text: str) -> str:
        """Suggest appropriate job title based on resume content"""
        # First check current/recent job titles
        experiences = self._extract_experience(text)
        if experiences and experiences[0]['title']:
            current_title = experiences[0]['title']
            
            # Match against standard titles
            if fuzz and process:
                matches = process.extract(current_title, self.job_titles_database, scorer=fuzz.ratio, limit=1)
                if matches and matches[0][1] > 80:  # 80% similarity threshold
                    return matches[0][0]
            
            return current_title
        
        # Fallback: Analyze skills and experience
        skills = self._extract_skills(text)
        years_exp = self._calculate_total_experience(text)
        
        # Determine seniority
        if years_exp < 2:
            seniority = "Junior"
        elif years_exp < 5:
            seniority = ""
        elif years_exp < 8:
            seniority = "Senior"
        else:
            seniority = "Lead"
        
        # Determine role based on skills
        if any(skill in skills for skill in ['React', 'Angular', 'Vue', 'Frontend']):
            role = "Frontend Developer"
        elif any(skill in skills for skill in ['Node.js', 'Django', 'Spring', 'Backend']):
            role = "Backend Developer"
        elif any(skill in skills for skill in ['Full Stack', 'MEAN', 'MERN']):
            role = "Full Stack Developer"
        elif any(skill in skills for skill in ['Machine Learning', 'Deep Learning', 'Data Science']):
            role = "Data Scientist"
        elif any(skill in skills for skill in ['AWS', 'Azure', 'DevOps', 'Kubernetes']):
            role = "DevOps Engineer"
        else:
            role = "Software Engineer"
        
        return f"{seniority} {role}".strip()
    
    def _parse_date_range(self, date_str: str) -> Tuple[Optional[str], Optional[str]]:
        """Parse date range from string"""
        # Handle "Present" or "Current"
        date_str = re.sub(r'\b(present|current|now)\b', datetime.now().strftime('%Y-%m'), date_str, flags=re.IGNORECASE)
        
        # Try to find two dates
        date_pattern = r'(\w+\.?\s*\d{4}|\d{1,2}[-/]\d{4}|\d{4})'
        dates = re.findall(date_pattern, date_str)
        
        if len(dates) >= 2:
            try:
                start = date_parser.parse(dates[0], fuzzy=True).strftime('%Y-%m')
                end = date_parser.parse(dates[1], fuzzy=True).strftime('%Y-%m')
                return start, end
            except:
                pass
        elif len(dates) == 1:
            try:
                start = date_parser.parse(dates[0], fuzzy=True).strftime('%Y-%m')
                return start, None
            except:
                pass
        
        return None, None
    
    def _calculate_duration(self, start_date: Optional[str], end_date: Optional[str]) -> str:
        """Calculate duration between dates"""
        if not start_date:
            return ""
        
        try:
            start = datetime.strptime(start_date, '%Y-%m')
            end = datetime.strptime(end_date, '%Y-%m') if end_date else datetime.now()
            
            delta = relativedelta(end, start)
            years = delta.years
            months = delta.months
            
            if years and months:
                return f"{years} years {months} months"
            elif years:
                return f"{years} years"
            elif months:
                return f"{months} months"
            else:
                return "Less than a month"
        except:
            return ""
    
    def _extract_bullet_points(self, text: str) -> List[str]:
        """Extract bullet points from text"""
        points = []
        
        # Common bullet indicators
        bullet_pattern = r'(?:^|\n)\s*[•·\-*]\s*([^\n•·\-*]+)'
        matches = re.findall(bullet_pattern, text)
        
        for match in matches:
            point = match.strip()
            if 10 < len(point) < 300:  # Reasonable length
                points.append(point)
        
        return points
    
    def _extract_field_of_study(self, degree_text: str) -> str:
        """Extract field of study from degree text"""
        # Common fields
        fields = [
            'Computer Science', 'Information Technology', 'Software Engineering',
            'Electrical Engineering', 'Electronics', 'Mechanical Engineering',
            'Data Science', 'Mathematics', 'Physics', 'Business Administration',
            'Finance', 'Marketing', 'Psychology', 'Economics'
        ]
        
        degree_lower = degree_text.lower()
        for field in fields:
            if field.lower() in degree_lower:
                return field
        
        # Extract text after "in" or "of"
        field_pattern = r'(?:in|of)\s+([A-Za-z\s]+?)(?:from|at|,|\n|$)'
        match = re.search(field_pattern, degree_text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        
        return ""


# Convenience function for direct usage
def parse_resume_file(file_path: str) -> Dict[str, Any]:
    """Parse a resume file and return structured data"""
    parser = EnhancedResumeParser()
    text = parser.extract_text_from_file(file_path)
    return parser.parse_resume(text, os.path.basename(file_path))


def parse_resume_text(text: str, filename: str = "") -> Dict[str, Any]:
    """Parse resume text and return structured data"""
    parser = EnhancedResumeParser()
    return parser.parse_resume(text, filename)
