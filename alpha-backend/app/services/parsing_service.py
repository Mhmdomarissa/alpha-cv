"""
Parsing Service - Consolidated Document Processing
Handles ALL text extraction from various file formats with PII removal.
Single responsibility: Extract clean text from documents.
"""
import logging
import os
import re
from typing import Dict, Any, Optional
from pathlib import Path
import tempfile
import shutil

# Document processing libraries
import fitz  # PyMuPDF for PDF
import docx  # python-docx for DOCX
import pytesseract  # OCR for images
from PIL import Image
from io import BytesIO

logger = logging.getLogger(__name__)

class ParsingService:
    """
    Consolidated service for all document text extraction and processing.
    Combines PDF, DOCX, OCR, and PII removal into a single clean interface.
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        """Initialize the parsing service."""
        self.logger = logger
        self.logger.info("🔧 ParsingService initialized")
    
    def process_document(self, file_path: str, doc_type: str = "auto") -> Dict[str, Any]:
        """
        Main processing function - calls appropriate extraction method.
        
        Args:
            file_path: Path to the document file
            doc_type: Type of document ("cv", "jd", or "auto" for detection)
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            # Validate file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_size = os.path.getsize(file_path)
            if file_size > self.MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size} bytes (max: {self.MAX_FILE_SIZE})")
            
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.SUPPORTED_EXTENSIONS:
                raise ValueError(f"Unsupported file extension: {file_ext}")
            
            self.logger.info(f"📄 Processing document: {file_path} ({file_size:,} bytes)")
            
            # Extract text based on file type
            if file_ext == '.pdf':
                raw_text = self.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                raw_text = self.extract_text_from_docx(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                raw_text = self.extract_with_ocr(file_path)
            elif file_ext == '.txt':
                raw_text = self._extract_text_from_txt(file_path)
            else:
                raise ValueError(f"No extraction method for: {file_ext}")
            
            # Clean and validate extracted text
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError(f"Insufficient text extracted: {len(raw_text)} characters")
            
            # Remove PII data
            clean_text, extracted_pii = self.remove_pii_data(raw_text)
            
            result = {
        "raw_text": raw_text,
        "clean_text": clean_text,
        "extracted_pii": extracted_pii,  # New field
        "file_path": file_path,
        "file_size": file_size,
        "file_extension": file_ext,
        "character_count": len(clean_text),
        "processing_status": "success"
    }

            self.logger.info(f"✅ Document processed successfully: {len(clean_text):,} characters")
            return result
            
        except Exception as e:
            self.logger.error(f"❌ Document processing failed: {str(e)}")
            raise Exception(f"Document processing failed: {str(e)}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF files using PyMuPDF.
        Combines all PDF extraction logic into one optimized function.
        """
        try:
            self.logger.info(f"📖 Extracting text from PDF: {file_path}")
            
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                if page_text.strip():  # Only add pages with content
                    text_parts.append(page_text)
                    
                # Also extract text from images in PDF if needed
                try:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        # Extract image and perform OCR if text is sparse
                        if len(page_text.strip()) < 100:  # Fallback to OCR for image-heavy pages
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_data = pix.tobytes("png")
                                ocr_text = self._perform_ocr_on_bytes(img_data)
                                if ocr_text.strip():
                                    text_parts.append(f"\n[OCR from page {page_num + 1}]: {ocr_text}")
                            pix = None
                except Exception as ocr_error:
                    self.logger.warning(f"OCR fallback failed for page {page_num + 1}: {str(ocr_error)}")
            
            doc.close()
            
            extracted_text = "\n".join(text_parts).strip()
            
            if not extracted_text:
                raise ValueError("No text could be extracted from PDF")
                
            self.logger.info(f"✅ PDF extraction successful: {len(extracted_text):,} characters")
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"❌ PDF extraction failed: {str(e)}")
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX files using python-docx.
        Includes paragraphs and tables.
        """
        try:
            self.logger.info(f"📝 Extracting text from DOCX: {file_path}")
            
            document = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in document.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append("\n--- Table Data ---")
                    text_parts.extend(table_text)
            
            extracted_text = "\n".join(text_parts).strip()
            
            if not extracted_text:
                raise ValueError("No text could be extracted from DOCX file")
                
            self.logger.info(f"✅ DOCX extraction successful: {len(extracted_text):,} characters")
            return extracted_text
            
        except Exception as e:
            self.logger.error(f"❌ DOCX extraction failed: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def extract_with_ocr(self, file_path: str) -> str:
        """
        OCR fallback for image-based documents using pytesseract.
        """
        try:
            self.logger.info(f"🔍 Performing OCR on image: {file_path}")
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, config='--psm 1')
            
            if not text.strip():
                raise ValueError("No text could be extracted via OCR")
                
            self.logger.info(f"✅ OCR extraction successful: {len(text):,} characters")
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"❌ OCR extraction failed: {str(e)}")
            raise Exception(f"OCR extraction failed: {str(e)}")
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text files with encoding detection."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
                return text.strip()
                
        except Exception as e:
            raise Exception(f"Failed to read text file: {str(e)}")
    
    def _perform_ocr_on_bytes(self, image_bytes: bytes) -> str:
        """Perform OCR on image bytes."""
        try:
            image = Image.open(BytesIO(image_bytes))
            text = pytesseract.image_to_string(image, config='--psm 1')
            return text.strip()
        except Exception:
            return ""
    
    def remove_pii_data(self, text: str) -> tuple[str, dict]:
        """
        Remove PII from text and return cleaned text plus extracted PII.
        Returns tuple: (clean_text, extracted_pii)
        """
        extracted_pii = {"email": [], "phone": []}
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            extracted_pii["email"] = list(set(emails))
        
        # Extract phone numbers
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'\+?[0-9]{1,4}[-.\s]?\(?[0-9]{1,4}\)?[-.\s]?[0-9]{1,4}[-.\s]?[0-9]{1,9}',
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        ]
        phones = []
        for pattern in phone_patterns:
            phones.extend(re.findall(pattern, text))
        if phones:
            extracted_pii["phone"] = list(set(phones))
        
        # Remove PII from text
        clean_text = text
        clean_text = re.sub(email_pattern, '[EMAIL]', clean_text)
        for pattern in phone_patterns:
            clean_text = re.sub(pattern, '[PHONE]', clean_text)
        
        # Clean up multiple whitespaces
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        return clean_text, extracted_pii

# Global instance
_parsing_service: Optional[ParsingService] = None

def get_parsing_service() -> ParsingService:
    """Get global parsing service instance."""
    global _parsing_service
    if _parsing_service is None:
        _parsing_service = ParsingService()
    return _parsing_service