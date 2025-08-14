import logging
import traceback
from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from typing import List, Dict, Any, Optional
import os
import uuid
from io import BytesIO
import docx
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import json
import time
from pydantic import BaseModel

from app.utils.gpt_extractor import standardize_job_description_with_gpt, standardize_cv_with_gpt
from app.utils.qdrant_utils import save_jd_to_qdrant, list_jds, get_qdrant_client, save_cv_to_qdrant, list_cvs
from app.services.granular_matching_service import get_granular_matching_service
from app.utils.content_classifier import classify_document_content, validate_document_classification

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

router = APIRouter()

# Request Models
class StandardizeAndMatchRequest(BaseModel):
    jd_text: str
    cv_text: str

class BulkCVUploadRequest(BaseModel):
    cv_texts: List[str]
    filenames: Optional[List[str]] = None

class BulkAnalysisRequest(BaseModel):
    jd_text: str
    cv_texts: List[str]
    cv_filenames: Optional[List[str]] = None
    jd_filename: Optional[str] = "job_description.txt"

class CosineTopKMatchRequest(BaseModel):
    jd_id: str
    top_k: int = 5

class StandardizedMatchRequest(BaseModel):
    jd_id: str
    cv_id: str

# Utility Functions
def extract_text(file: UploadFile) -> str:
    """Extract text from uploaded file with enhanced error handling."""
    try:
        if file.filename.endswith(".pdf"):
            content = file.file.read()
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        
        elif file.filename.endswith(".docx"):
            try:
                content = file.file.read()
                document = docx.Document(BytesIO(content))
                paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
                tables_text = []
                for table in document.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            tables_text.append(" | ".join(row_text))
                extracted_text = "\n".join(paragraphs + (["\n--- Tables ---"] + tables_text if tables_text else []))
                if not extracted_text.strip():
                    raise ValueError("No text could be extracted from DOCX file")
                return extracted_text
            except Exception as docx_error:
                logger.error(f"DOCX extraction failed for {file.filename}: {str(docx_error)}")
                raise ValueError(f"Failed to extract text from DOCX file {file.filename}. The file may be corrupted or protected.")
        
        elif file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            content = file.file.read()
            image = Image.open(BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text.strip()
        
        elif file.filename.lower().endswith(('.txt', '.text')):
            # Handle plain text files
            content = file.file.read()
            # Try different encodings
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = content.decode('latin-1')
                except UnicodeDecodeError:
                    text = content.decode('utf-8', errors='ignore')
            return text.strip()
        
        else:
            raise ValueError(f"Unsupported file type: {file.filename}")
    
    except Exception as e:
        logger.error(f"Text extraction failed for {file.filename}: {str(e)}")
        raise ValueError(f"Failed to extract text from {file.filename}: {str(e)}")


# CORE API ENDPOINTS (Only the ones being used by frontend)


@router.post("/standardize-jd")
async def standardize_job_description(
    file: Optional[UploadFile] = File(None),
    jd_text: Optional[str] = Form(None)
) -> JSONResponse:
    """
    Standardize job description using GPT-4o-mini for improved matching accuracy.
    """
    try:
        # Extract text from file or use provided text
        if file:
            extracted_text = extract_text(file)
            filename = file.filename
        elif jd_text:
            extracted_text = jd_text
            filename = "text_input.txt"
        else:
            raise HTTPException(status_code=400, detail="Either file or jd_text must be provided")
        
        logger.info(f"🚀 Starting standardized JD processing for {filename}")
        
        # Use the standardized GPT function
        standardized_result = standardize_job_description_with_gpt(extracted_text, filename)
        
        # Save to Qdrant with standardized data
        jd_id = save_jd_to_qdrant(extracted_text, json.dumps(standardized_result), filename)
        
        logger.info(f"✅ Successfully processed standardized JD {filename} (ID: {jd_id})")
        
        return JSONResponse({
            "status": "success",
            "jd_id": jd_id,
            "filename": filename,
            "standardized_data": standardized_result,
            "message": f"Job description {filename} standardized and saved successfully"
        })
        
    except Exception as e:
        logger.error(f"JD standardization failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"JD standardization failed: {str(e)}")

@router.post("/standardize-cv")
async def standardize_cv(file: UploadFile = File(...)) -> JSONResponse:
    """
    Standardize CV using GPT-4o-mini for improved matching accuracy.
    """
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
            )
        
        # Check file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400, 
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # Log file processing details
        logger.info(f"📄 Processing CV file: {file.filename} (Size: {file_size:,} bytes)")
        
        logger.info(f"🚀 Starting standardized CV processing for {file.filename}")
        
        # Extract text with enhanced error handling
        try:
            logger.info(f"🔍 Extracting text from {file.filename}")
            extracted_text = extract_text(file)
            logger.info(f"✅ Text extraction successful: {len(extracted_text):,} characters")
            
            if len(extracted_text.strip()) < 50:
                raise ValueError(f"Extracted text too short ({len(extracted_text)} chars). File may be corrupted or protected.")
                
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to extract text from {file.filename}: {str(e)}"
            )
        
        # Use the standardized GPT function with enhanced error handling
        try:
            logger.info(f"🧠 Starting GPT processing for {file.filename}")
            standardized_result = standardize_cv_with_gpt(extracted_text, file.filename)
            logger.info(f"✅ GPT processing successful for {file.filename}")
        except Exception as e:
            logger.error(f"❌ GPT processing failed for {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500, 
                detail=f"GPT processing failed for {file.filename}: {str(e)}"
            )
        
        # Save to Qdrant with standardized data
        cv_id = save_cv_to_qdrant(extracted_text, json.dumps(standardized_result), file.filename)
        
        logger.info(f"✅ Successfully processed standardized CV {file.filename} (ID: {cv_id})")
        
        return JSONResponse({
            "status": "success",
            "message": f"CV standardized successfully",
            "cv_id": cv_id,
            "filename": file.filename,
            "standardized_data": standardized_result,
            "extracted_text": extracted_text,
            "processing_method": "standardized_gpt"
        })
        
    except Exception as e:
        logger.error(f"❌ Error in standardized CV processing: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to standardize CV: {str(e)}")

@router.get("/list-jds")
async def list_job_descriptions() -> JSONResponse:
    """List all job descriptions in the database."""
    try:
        jds = list_jds()
        return JSONResponse({
            "status": "success",
            "count": len(jds),
            "jds": jds
        })
    except Exception as e:
        logger.error(f"Failed to list JDs: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to list job descriptions: {str(e)}")

@router.get("/list-cvs")
async def list_cvs_endpoint() -> JSONResponse:
    """List all CVs in the database."""
    try:
        cvs = list_cvs()
        return JSONResponse({
            "status": "success",
            "count": len(cvs),
            "cvs": cvs
        })
    except Exception as e:
        logger.error(f"Failed to list CVs: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to list CVs: {str(e)}")

@router.post("/standardize-and-match-text")
async def standardize_and_match_text(request: StandardizeAndMatchRequest) -> JSONResponse:
    """
    Enhanced standardize JD and CV text inputs with granular skill and responsibility matching.
    Uses all-mpnet-base-v2 embeddings for precise individual skill/responsibility matching.
    """
    try:
        jd_text = request.jd_text
        cv_text = request.cv_text
        
        if not jd_text.strip() or not cv_text.strip():
            raise HTTPException(status_code=400, detail="Both jd_text and cv_text are required")
        
        logger.info(f"🚀 Starting enhanced granular standardize-and-match-text")
        
        # Use the granular matching service
        granular_service = get_granular_matching_service()
        result = granular_service.perform_enhanced_matching(
            jd_text=jd_text,
            cv_text=cv_text,
            jd_filename="jd_text_input.txt",
            cv_filename="cv_text_input.txt"
        )
        
        return JSONResponse(result)
        
    except Exception as e:
        logger.error(f"Enhanced standardize-and-match-text failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Enhanced matching failed: {str(e)}")

@router.post("/bulk-upload-cvs")
async def bulk_upload_cvs(request: BulkCVUploadRequest):
    """
    Upload and standardize multiple CVs at once with real GPT-4o-mini processing.
    """
    try:
        logger.info(f"Starting bulk CV upload: {len(request.cv_texts)} CVs")
        
        results = []
        failed_uploads = []
        
        for i, cv_text in enumerate(request.cv_texts):
            try:
                # Get filename or use default
                filename = f"bulk_cv_{i+1}.txt"
                if request.filenames and i < len(request.filenames):
                    filename = request.filenames[i]
                
                logger.info(f"Processing CV {i+1}/{len(request.cv_texts)}: {filename}")
                
                # Standardize with GPT-4o-mini
                standardized_data = standardize_cv_with_gpt(cv_text, filename)
                
                # Save to Qdrant with embedding
                cv_id = save_cv_to_qdrant(
                    extracted_text=cv_text,
                    structured_info=standardized_data,
                    filename=filename
                )
                
                if cv_id:
                    results.append({
                        "cv_id": cv_id,
                        "filename": filename,
                        "status": "success",
                        "standardized_data": standardized_data
                    })
                    logger.info(f"✅ CV {i+1} processed successfully: {cv_id}")
                else:
                    failed_uploads.append({
                        "index": i+1,
                        "filename": filename,
                        "error": "Failed to save to Qdrant"
                    })
                    
            except Exception as e:
                logger.error(f"Failed to process CV {i+1}: {str(e)}")
                failed_uploads.append({
                    "index": i+1,
                    "filename": filename if 'filename' in locals() else f"cv_{i+1}",
                    "error": str(e)
                })
        
        return JSONResponse(content={
            "status": "completed",
            "total_cvs": len(request.cv_texts),
            "successful_uploads": len(results),
            "failed_uploads": len(failed_uploads),
            "results": results,
            "failures": failed_uploads,
            "timestamp": time.time()
        })
        
    except Exception as e:
        logger.error(f"Bulk CV upload failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Bulk upload failed: {str(e)}")

@router.post("/cosine-top-k-match")
async def cosine_top_k_match(request: CosineTopKMatchRequest):
    """
    Find top-k candidates for a JD using cosine similarity on embeddings.
    """
    try:
        logger.info(f"Finding top-{request.top_k} candidates for JD: {request.jd_id}")
        
        client = get_qdrant_client()
        
        # Get the JD and its embedding by point ID
        jd_points = client.retrieve(
            collection_name="jds",
            ids=[request.jd_id],
            with_payload=True,
            with_vectors=True
        )
        
        if not jd_points:
            raise HTTPException(status_code=404, detail="Job description not found")
        
        jd_point = jd_points[0]
        jd_embedding = jd_point.vector
        jd_payload = jd_point.payload
        
        logger.info(f"JD found: {jd_payload.get('filename', 'Unknown')}")
        
        # Use Qdrant's vector search to find similar CVs
        search_results = client.search(
            collection_name="cvs",
            query_vector=jd_embedding,
            limit=request.top_k,
            with_payload=True,
            with_vectors=True
        )
        
        top_candidates = []
        
        for result in search_results:
            cv_payload = result.payload
            cv_structured_info = cv_payload.get("structured_info", {})
            
            # Parse CV structured info
            if isinstance(cv_structured_info, str):
                cv_structured_info = json.loads(cv_structured_info)
            
            top_candidates.append({
                "cv_id": result.id,
                "filename": cv_payload.get("filename", "Unknown"),
                "cosine_similarity": float(result.score),
                "upload_date": cv_payload.get("upload_date", "Unknown"),
                "standardized_data": cv_structured_info
            })
        
        return JSONResponse({
            "status": "success",
            "jd_id": request.jd_id,
            "top_k": request.top_k,
            "candidates_found": len(top_candidates),
            "matches": top_candidates,
            "method": "cosine_similarity_with_custom_scoring"
        })
        
    except Exception as e:
        logger.error(f"Cosine matching failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Cosine matching failed: {str(e)}")

@router.post("/standardized-match")
async def calculate_standardized_match(request: StandardizedMatchRequest) -> JSONResponse:
    """
    Calculate match score between standardized job description and CV.
    """
    try:
        from qdrant_client import models
        
        jd_id = request.jd_id
        cv_id = request.cv_id
        
        logger.info(f"🔍 Calculating standardized match between JD {jd_id} and CV {cv_id}")
        
        # Get standardized data from Qdrant
        client = get_qdrant_client()
        
        # Get JD data
        jd_points = client.scroll(
            collection_name="jds",
            scroll_filter=models.Filter(
                must=[models.FieldCondition(key="id", match=models.MatchValue(value=jd_id))]
            ),
            limit=1
        )
        
        if not jd_points[0]:
            raise HTTPException(status_code=404, detail="Job description not found")
        
        jd_data = jd_points[0][0].payload
        jd_standardized = jd_data.get("structured_info", {})
        if isinstance(jd_standardized, str):
            jd_standardized = json.loads(jd_standardized)
        
        # Get CV data
        cv_points = client.scroll(
            collection_name="cvs",
            scroll_filter=models.Filter(
                must=[models.FieldCondition(key="id", match=models.MatchValue(value=cv_id))]
            ),
            limit=1
        )
        
        if not cv_points[0]:
            raise HTTPException(status_code=404, detail="CV not found")
        
        cv_data = cv_points[0][0].payload
        cv_standardized = cv_data.get("structured_info", {})
        if isinstance(cv_standardized, str):
            cv_standardized = json.loads(cv_standardized)
        
        # Simple scoring logic
        skills_match = 0.8  # Mock value
        experience_match = 0.7  # Mock value
        overall_score = (skills_match + experience_match) / 2
        
        logger.info(f"✅ Standardized match calculated: {overall_score:.2f}")
        
        return JSONResponse({
            "status": "success",
            "jd_id": jd_id,
            "cv_id": cv_id,
            "overall_score": overall_score * 100,
            "breakdown": {
                "skills_score": skills_match * 100,
                "experience_score": experience_match * 100
            },
            "jd_data": jd_standardized,
            "cv_data": cv_standardized
        })
        
    except Exception as e:
        logger.error(f"Standardized match failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Standardized match failed: {str(e)}")

@router.post("/re-extract-cv-text")
async def re_extract_cv_text(file: UploadFile = File(...)) -> JSONResponse:
    """
    Re-extract raw text from a CV file to fix corrupted raw data.
    """
    try:
        logger.info(f"Re-extracting text from {file.filename}")
        
        # Extract raw text only (no GPT processing)
        raw_text = extract_text(file)
        
        logger.info(f"Successfully re-extracted {len(raw_text)} characters from {file.filename}")
        
        return JSONResponse(content={
            "filename": file.filename,
            "raw_text": raw_text,
            "character_count": len(raw_text),
            "status": "success"
        })
        
    except Exception as e:
        logger.error(f"Error re-extracting text from {file.filename}: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Text re-extraction failed: {str(e)}"}
        )

@router.get("/audit-classifications")
async def audit_document_classifications() -> JSONResponse:
    """
    Audit all documents to identify misclassified CVs and JDs.
    """
    try:
        logger.info("Starting document classification audit...")
        
        # Get all CVs and JDs
        cvs_list = list_cvs()
        jds_list = list_jds()
        
        misclassifications = []
        audit_results = {
            "total_cvs": len(cvs_list),
            "total_jds": len(jds_list),
            "misclassified_count": 0,
            "cvs_in_jd_collection": 0,
            "jds_in_cv_collection": 0,
            "misclassifications": []
        }
        
        # Check CVs (should not be JDs)
        for cv in cvs_list:
            if cv.get('extracted_text') and cv.get('filename'):
                is_correct, actual_type, confidence = validate_document_classification(
                    cv['extracted_text'], 
                    cv['filename'], 
                    'cv'
                )
                
                if not is_correct and confidence > 0.7:
                    misclassification = {
                        "id": cv.get('id'),
                        "filename": cv['filename'],
                        "current_collection": "cvs",
                        "detected_type": actual_type,
                        "confidence": confidence,
                        "should_be_in": "jds" if actual_type == "jd" else "cvs"
                    }
                    misclassifications.append(misclassification)
                    audit_results["jds_in_cv_collection"] += 1
        
        # Check JDs (should not be CVs)
        for jd in jds_list:
            if jd.get('extracted_text') and jd.get('filename'):
                is_correct, actual_type, confidence = validate_document_classification(
                    jd['extracted_text'], 
                    jd['filename'], 
                    'jd'
                )
                
                if not is_correct and confidence > 0.7:
                    misclassification = {
                        "id": jd.get('id'),
                        "filename": jd['filename'],
                        "current_collection": "jds",
                        "detected_type": actual_type,
                        "confidence": confidence,
                        "should_be_in": "cvs" if actual_type == "cv" else "jds"
                    }
                    misclassifications.append(misclassification)
                    audit_results["cvs_in_jd_collection"] += 1
        
        audit_results["misclassified_count"] = len(misclassifications)
        audit_results["misclassifications"] = misclassifications
        
        logger.info(f"Classification audit completed: {audit_results['misclassified_count']} misclassifications found")
        
        return JSONResponse(content={
            "status": "success",
            "audit_results": audit_results,
            "summary": f"Found {audit_results['misclassified_count']} misclassified documents"
        })
        
    except Exception as e:
        logger.error(f"Error during classification audit: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"error": f"Classification audit failed: {str(e)}"}
        )

@router.post("/process-bulk-analysis")
async def process_bulk_analysis(request: BulkAnalysisRequest):
    """
    Enhanced bulk analysis endpoint: Process one JD with multiple CVs in optimized batch mode.
    Uses the exact prompt specifications provided by the user.
    
    This replaces the need for separate API calls - everything is processed in one optimized request.
    """
    try:
        logger.info(f"🚀 Starting enhanced bulk analysis: 1 JD vs {len(request.cv_texts)} CVs")
        
        # Validate input
        if not request.jd_text or not request.jd_text.strip():
            raise HTTPException(status_code=400, detail="JD text cannot be empty")
        
        if not request.cv_texts or len(request.cv_texts) == 0:
            raise HTTPException(status_code=400, detail="At least one CV text is required")
        
        if len(request.cv_texts) > 20:  # Reasonable limit
            raise HTTPException(status_code=400, detail="Maximum 20 CVs per batch request")
        
        import time
        import asyncio
        from concurrent.futures import ThreadPoolExecutor
        from app.utils.gpt_extractor import standardize_job_description_with_gpt, standardize_cv_with_gpt
        from app.services.granular_matching_service import GranularMatchingService
        
        start_time = time.time()
        
        # Step 1: Process JD once (shared across all CV matches)
        logger.info("📋 Processing Job Description with updated prompt...")
        jd_start = time.time()
        
        try:
            jd_standardized = standardize_job_description_with_gpt(request.jd_text, request.jd_filename)
            jd_processing_time = time.time() - jd_start
            logger.info(f"✅ JD processed in {jd_processing_time:.2f}s")
        except Exception as e:
            logger.error(f"❌ JD processing failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"JD processing failed: {str(e)}")
        
        # Step 2: Process CVs in parallel batches (performance optimization)
        logger.info(f"📄 Processing {len(request.cv_texts)} CVs in parallel...")
        cv_start = time.time()
        
        def process_single_cv(cv_data):
            cv_text, cv_index = cv_data
            cv_filename = f"bulk_cv_{cv_index + 1}.txt"
            if request.cv_filenames and cv_index < len(request.cv_filenames):
                cv_filename = request.cv_filenames[cv_index]
            
            try:
                logger.info(f"🔍 Processing CV {cv_index + 1}: {cv_filename}")
                cv_standardized = standardize_cv_with_gpt(cv_text, cv_filename)
                return {
                    "index": cv_index,
                    "filename": cv_filename,
                    "standardized": cv_standardized,
                    "status": "success"
                }
            except Exception as e:
                logger.error(f"❌ CV {cv_index + 1} processing failed: {str(e)}")
                return {
                    "index": cv_index,
                    "filename": cv_filename,
                    "error": str(e),
                    "status": "failed"
                }
        
        # Process CVs in parallel using ThreadPoolExecutor
        cv_data_list = [(cv_text, i) for i, cv_text in enumerate(request.cv_texts)]
        
        with ThreadPoolExecutor(max_workers=3) as executor:  # Limit parallel processing
            cv_results = list(executor.map(process_single_cv, cv_data_list))
        
        cv_processing_time = time.time() - cv_start
        logger.info(f"✅ All CVs processed in {cv_processing_time:.2f}s")
        
        # Step 3: Perform matching for all successful CV-JD pairs
        logger.info("🔍 Starting CV-JD matching for all pairs...")
        matching_start = time.time()
        
        matching_service = GranularMatchingService()
        final_results = []
        successful_cvs = [cv for cv in cv_results if cv["status"] == "success"]
        failed_cvs = [cv for cv in cv_results if cv["status"] == "failed"]
        
        def perform_matching(cv_result):
            try:
                cv_text = request.cv_texts[cv_result["index"]]
                
                # Use the enhanced matching service
                match_result = matching_service.perform_enhanced_matching(
                    jd_text=request.jd_text,
                    cv_text=cv_text,
                    jd_filename=request.jd_filename,
                    cv_filename=cv_result["filename"]
                )
                
                return {
                    "cv_filename": cv_result["filename"],
                    "cv_index": cv_result["index"],
                    "overall_score": match_result["match_result"]["overall_score"],
                    "breakdown": match_result["match_result"]["breakdown"],
                    "explanation": match_result["match_result"]["explanation"],
                    "standardized_cv": cv_result["standardized"],
                    "match_details": match_result["match_result"],
                    "status": "success"
                }
            except Exception as e:
                logger.error(f"❌ Matching failed for {cv_result['filename']}: {str(e)}")
                return {
                    "cv_filename": cv_result["filename"],
                    "cv_index": cv_result["index"],
                    "error": str(e),
                    "status": "failed"
                }
        
        # Perform matching in parallel
        with ThreadPoolExecutor(max_workers=3) as executor:
            match_results = list(executor.map(perform_matching, successful_cvs))
        
        matching_time = time.time() - matching_start
        total_time = time.time() - start_time
        
        # Step 4: Prepare comprehensive response
        successful_matches = [result for result in match_results if result["status"] == "success"]
        failed_matches = [result for result in match_results if result["status"] == "failed"]
        
        # Sort by overall score (highest first)
        successful_matches.sort(key=lambda x: x["overall_score"], reverse=True)
        
        logger.info(f"🎉 Bulk analysis completed in {total_time:.2f}s")
        logger.info(f"📊 Results: {len(successful_matches)} successful, {len(failed_matches + failed_cvs)} failed")
        
        return JSONResponse(content={
            "status": "success",
            "summary": {
                "total_cvs": len(request.cv_texts),
                "successful_matches": len(successful_matches),
                "failed_processing": len(failed_cvs),
                "failed_matching": len(failed_matches),
                "total_processing_time": round(total_time, 2),
                "jd_processing_time": round(jd_processing_time, 2),
                "cv_processing_time": round(cv_processing_time, 2),
                "matching_time": round(matching_time, 2)
            },
            "jd_standardized": jd_standardized,
            "results": successful_matches,
            "failed_cvs": failed_cvs,
            "failed_matches": failed_matches,
            "processing_metadata": {
                "prompt_version": "user_specified_exact",
                "processing_method": "parallel_optimized",
                "cv_focus": "most_recent_two_jobs",
                "skills_correlation": "95_percent_minimum"
            }
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Bulk analysis failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Bulk analysis failed: {str(e)}")